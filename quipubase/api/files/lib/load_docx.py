import base64c as base64
import typing as tp
from dataclasses import dataclass

import typing_extensions as tpe
from docx import Document

from ._base import Artifact
from .utils import get_logger

logger = get_logger(__name__)


@dataclass
class DocxLoader(Artifact):
    ref: tpe.Annotated[
        str,
        tpe.Doc(
            """
		This `ref` can represent one out of three things:

		- An HTTP URL.
		- A file path (temporary or not) within the local filesystem.
		- A text file content.
		"""
        ),
    ]

    def extract(self) -> tp.Generator[str, None, None]:
        """
        Extract text and images from Word documents.

        Yields:
            Text content and image data as HTML
        """
        file_path = self.retrieve()

        try:
            # Open the DOCX file
            doc = Document(file_path.as_posix())

            # Process document properties
            prop_html: list[str] = []

            if hasattr(doc.core_properties, "title") and doc.core_properties.title:
                prop_html.append(f"<h1>{doc.core_properties.title}</h1>")

            if hasattr(doc.core_properties, "author") and doc.core_properties.author:
                prop_html.append(
                    f"<p><strong>Author:</strong> {doc.core_properties.author}</p>"
                )

            if prop_html:
                yield "\n".join(prop_html)

            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    # Get paragraph style information
                    style_name = paragraph.style.name if paragraph.style else "Normal"

                    # Handle headings
                    if style_name and style_name.startswith("Heading"):
                        heading_level = (
                            style_name.split(" ")[-1]
                            if " " in style_name
                            else style_name[-1]
                        )
                        try:
                            h_level = int(heading_level)
                            if 1 <= h_level <= 6:
                                yield f"<h{h_level}>{paragraph.text}</h{h_level}>"
                                continue
                        except (ValueError, TypeError) as e:
                            logger.error(
                                f"Invalid heading level: {heading_level}. Skipping paragraph."
                            )
                            logger.error("Error: %s - %s", e.__class__.__name__, e)

                    # Regular paragraph
                    yield f"<p>{paragraph.text}</p>"

                # Check for inline images in paragraph runs
                for run in paragraph.runs:
                    if not run.text:  # Empty run might contain an image
                        try:
                            # Attempt to extract images from XML elements
                            for inline in run._element.iter():  # type: ignore
                                if inline.tag.endswith("inline"):  # type: ignore
                                    for pic in inline.iter():  # type: ignore
                                        if pic.tag.endswith("blip"):  # type: ignore
                                            try:
                                                image_id = pic.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")  # type: ignore
                                                if image_id:
                                                    image_part = run.part.related_parts[
                                                        image_id
                                                    ]
                                                    img_data = base64.b64encode(
                                                        image_part.blob
                                                    ).decode()
                                                    # Try to determine image type from the part content type
                                                    content_type = (
                                                        image_part.content_type
                                                    )
                                                    img_type = (
                                                        content_type.split("/")[-1]
                                                        if "/" in content_type
                                                        else "png"
                                                    )
                                                    yield f'<img style="width: 24em;" src="data:image/{img_type};base64,{img_data}" />'
                                            except Exception as img_error:
                                                logger.error(
                                                    f"Error extracting image: {str(img_error)}"
                                                )
                                                yield f"<p>Error extracting image: {str(img_error)}</p>"
                        except Exception as e:
                            # If image extraction fails, move on
                            pass

            # Process tables
            for table in doc.tables:
                table_html = ["<table border='1' style='border-collapse: collapse;'>"]

                for row in table.rows:
                    table_html.append("<tr>")
                    for cell in row.cells:
                        table_html.append(f"<td>{cell.text}</td>")
                    table_html.append("</tr>")

                table_html.append("</table>")
                yield "\n".join(table_html)

        except Exception as e:
            # Handle errors gracefully
            yield f"Error processing Word document: {str(e)}"
